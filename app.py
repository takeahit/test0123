import pandas as pd
from rapidfuzz import fuzz, process
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
from pydocx import PyDocX  # .doc ファイルを扱うためのライブラリ
import streamlit as st
from PyPDF2 import PdfReader  # PDFからテキストを抽出するためのライブラリ
import re

# 制御文字やNULLバイトを除去する関数
def clean_strings(df):
    def clean_cell(value):
        if isinstance(value, str):
            # 制御文字とNULLバイトを除去
            return re.sub(r'[\x00-\x1F\x7F]', '', value)
        return value
    
    # データフレーム全体に適用
    return df.applymap(clean_cell)

# デバッグ用: 非互換文字を検出
def find_invalid_chars(df):
    invalid_rows = []
    for col in df.columns:
        for idx, value in df[col].items():
            if isinstance(value, str) and re.search(r'[\x00-\x1F\x7F]', value):
                invalid_rows.append((col, idx, value))
    return invalid_rows

# Excel ファイルを読み込む関数
def load_excel(file):
    df = pd.read_excel(file, engine="openpyxl")
    if df.columns.size < 1:
        raise ValueError("Excelファイルには少なくとも1列以上の用語が必要です。")
    return df

# Word、DOC または PDF ファイルからテキストを抽出する関数
def extract_text_from_file(file, file_type):
    if file_type == "docx":
        doc = Document(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_type == "doc":
        return PyDocX.to_text(file)
    elif file_type == "pdf":
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            page_text = page_text.replace("\n", " ").replace("\r", " ")
            page_text = " ".join(page_text.split())
            text += page_text + " "
        text = text.strip()
        return text
    else:
        return ""

# Fuzzy Matching を用いて類似語を検出する関数
def find_similar_terms(text, terms, threshold):
    words = text.split()
    detected_terms = []

    for word in words:
        matches = process.extract(word, terms, scorer=fuzz.partial_ratio, limit=10)
        for match in matches:
            if match[1] >= threshold and match[1] < 100:
                detected_terms.append((word, match[0], match[1]))

    return detected_terms

# 修正を適用して新しい Word ファイルを作成する関数
def create_corrected_word_file_with_formatting(original_text, corrections):
    doc = Document()
    for paragraph_text in original_text.split("\n"):
        paragraph = doc.add_paragraph()
        start_index = 0

        for incorrect, correct in corrections:
            while incorrect in paragraph_text[start_index:]:
                start_index = paragraph_text.find(incorrect, start_index)
                end_index = start_index + len(incorrect)
                paragraph.add_run(paragraph_text[:start_index])
                run = paragraph.add_run(correct)
                run.font.highlight_color = 6
                paragraph_text = paragraph_text[end_index:]
                start_index = 0
        paragraph.add_run(paragraph_text)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# 修正箇所を表にまとめる関数
def create_correction_table(detected):
    correction_table = pd.DataFrame(detected, columns=["原稿内の語", "類似する用語", "類似度"])
    return correction_table

# 正誤表を使用して修正を適用する関数
def apply_corrections_with_table(text, correction_df):
    corrections = []
    total_replacements = 0
    for _, row in correction_df.iterrows():
        incorrect, correct = row.iloc[0], row.iloc[1]
        max_replacements = text.count(incorrect)
        for _ in range(max_replacements):
            if incorrect in text:
                corrections.append((incorrect, correct))
                text = text.replace(incorrect, correct, 1)
                total_replacements += 1
            else:
                break
    return text, corrections, total_replacements

# 利用漢字表を使用して修正を適用する関数
def apply_kanji_table(text, kanji_df):
    corrections = []
    total_replacements = 0
    for _, row in kanji_df.iterrows():
        hiragana, kanji = row.iloc[0], row.iloc[1]
        max_replacements = text.count(hiragana)
        for _ in range(max_replacements):
            if hiragana in text:
                corrections.append((hiragana, kanji))
                text = text.replace(hiragana, kanji, 1)
                total_replacements += 1
            else:
                break
    return text, corrections, total_replacements

# Streamlit アプリケーション
st.markdown("<h1 style='text-align: center;'>南江堂用用語チェッカー（笑）</h1>", unsafe_allow_html=True)

st.write("以下のファイルを個別にアップロードしてください:")
word_file = st.file_uploader("原稿ファイル (Word, DOC, PDF):", type=["docx", "doc", "pdf"])
terms_file = st.file_uploader("用語集ファイル (A列に正しい用語を記載したExcel):", type=["xlsx"])
correction_file = st.file_uploader("正誤表ファイル (A列に誤った用語、B列に正しい用語を記載したExcel):", type=["xlsx"])
kanji_file = st.file_uploader("利用漢字表ファイル (A列にひらがな、B列に漢字を記載したExcel):", type=["xlsx"])

# アップロードファイルサイズの制限 (100MB以下)
if word_file and word_file.size > 100 * 1024 * 1024:
    st.error("原稿ファイルのサイズが大きすぎます（100MB以下にしてください）。")
    st.stop()

if terms_file and terms_file.size > 100 * 1024 * 1024:
    st.error("用語集ファイルのサイズが大きすぎます（100MB以下にしてください）。")
    st.stop()

if correction_file and correction_file.size > 100 * 1024 * 1024:
    st.error("正誤表ファイルのサイズが大きすぎます（100MB以下にしてください）。")
    st.stop()

if kanji_file and kanji_file.size > 100 * 1024 * 1024:
    st.error("利用漢字表ファイルのサイズが大きすぎます（100MB以下にしてください）。")
    st.stop()

if word_file and (terms_file or correction_file or kanji_file):
    file_type = word_file.name.split(".")[-1]
    original_text = extract_text_from_file(word_file, file_type)

    corrections = []

    if terms_file:
        try:
            terms_df = load_excel(terms_file)
            terms_df = clean_strings(terms_df)
            terms = terms_df.iloc[:, 0].dropna().astype(str).tolist()
            threshold = st.slider("類似度の閾値を設定してください (50-99):", min_value=50, max_value=99, value=65)
            detected = find_similar_terms(original_text, terms, threshold)

            if detected:
                st.success(f"類似語が{len(detected)}件検出されました！")
                correction_table = create_correction_table(detected)
                st.dataframe(correction_table)

                output = BytesIO()
                with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                    correction_table.to_excel(writer, index=False, sheet_name="修正箇所一覧")
                st.download_button(
                    label="修正箇所一覧をダウンロード",
                    data=output.getvalue(),
                    file_name="修正箇所一覧.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

        except Exception as e:
            st.error(f"用語集ファイルの処理中にエラーが発生しました: {e}")

    if correction_file:
        try:
            correction_df = load_excel(correction_file)
            correction_df = clean_strings(correction_df)
            invalid_chars = find_invalid_chars(correction_df)
            if invalid_chars:
                st.error(f"非互換文字が検出されました: {invalid_chars}")
            else:
                updated_text, corrections_from_table, total_replacements = apply_corrections_with_table(original_text, correction_df)
                corrections.extend(corrections_from_table)

                st.success(f"正誤表を適用し、{total_replacements}回の修正を行いました！")

                corrections_df = pd.DataFrame(corrections_from_table, columns=["誤った用語", "正しい用語"])
                st.dataframe(corrections_df)

                output = BytesIO()
                with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                    corrections_df.to_excel(writer, index=False, sheet_name="正誤表修正箇所")
                st.download_button(
                    label="正誤表修正箇所をダウンロード",
                    data=output.getvalue(),
                    file_name="正誤表修正箇所.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

                corrected_file = create_corrected_word_file_with_formatting(original_text, corrections_from_table)
                st.download_button(
                    label="正誤表修正済みファイルをダウンロード",
                    data=corrected_file.getvalue(),
                    file_name="正誤表修正済み.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        except Exception as e:
            st.error(f"正誤表ファイルの処理中にエラーが発生しました: {e}")

    if kanji_file:
        try:
            kanji_df = load_excel(kanji_file)
            kanji_df = clean_strings(kanji_df)
            updated_text, kanji_corrections, total_replacements = apply_kanji_table(original_text, kanji_df)
            corrections.extend(kanji_corrections)

            st.success(f"利用漢字表を適用し、{total_replacements}回の修正を行いました！")

            kanji_corrections_df = pd.DataFrame(kanji_corrections, columns=["ひらがな", "漢字"])
            st.dataframe(kanji_corrections_df)

            output = BytesIO()
            with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                kanji_corrections_df.to_excel(writer, index=False, sheet_name="漢字修正箇所")
            st.download_button(
                label="利用漢字表修正箇所をダウンロード",
                data=output.getvalue(),
                file_name="利用漢字表修正箇所.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

            corrected_file = create_corrected_word_file_with_formatting(original_text, kanji_corrections)
            st.download_button(
                label="利用漢字表修正済みファイルをダウンロード",
                data=corrected_file.getvalue(),
                file_name="利用漢字表修正済み.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        except Exception as e:
            st.error(f"利用漢字表ファイルの処理中にエラーが発生しました: {e}")

else:
    st.warning("原稿ファイルと、用語集、正誤表、利用漢字表のいずれかをアップロードしてください！")
